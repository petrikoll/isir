from __future__ import annotations

from datetime import datetime, timedelta
from datetime import date
from copy import deepcopy
from io import BytesIO
import json
import os
import shutil
import sqlite3
import sys
import tempfile
import threading
import unicodedata
from urllib.parse import quote, unquote
import zipfile

import calendar
import re
from pathlib import Path

from flask import Flask, abort, make_response, redirect, render_template, request, send_file, url_for
from flask_wtf import CSRFProtect
from docx import Document
from docx.shared import Inches, Pt
from openpyxl import load_workbook
from sqlalchemy.exc import IntegrityError

from ai_analysis import analyze_case_data_verification_job, analyze_case_latest_document_job, analyze_case_study_job
from app_settings import get_secret_key, has_gemini_api_key, set_gemini_api_key
from models import Client, DATABASE_PATH, SessionLocal, engine, init_db
from models import InsolvencyCase, InsolvencyDocument
from scheduler import check_all_clients, start_scheduler
from storage_paths import DOCUMENTS_DIR


def resource_path(relative_path: str) -> str:
    if getattr(sys, "frozen", False):
        return str(Path(sys._MEIPASS) / relative_path)
    return str(Path(__file__).resolve().parent / relative_path)


app = Flask(__name__, template_folder=resource_path("templates"))
app.config["SECRET_KEY"] = get_secret_key()
csrf = CSRFProtect(app)
init_db()
scheduler = start_scheduler()

PROJECT_FILTER_OPTIONS = ["SRSS II", "SRSS III", "CECH 2026+", "CECH 2025-"]
PROJECT_IMPORT_ALIASES = {
    "jpd ii": "SRSS II",
    "jpd 2": "SRSS II",
    "jlpd ii": "SRSS II",
    "jlpd 2": "SRSS II",
    "srss ii": "SRSS II",
    "srss 2": "SRSS II",
    "jpd iii": "SRSS III",
    "jpd 3": "SRSS III",
    "srs": "SRSS III",
    "srss iii": "SRSS III",
    "srss 3": "SRSS III",
    "cech 26+": "CECH 2026+",
    "cech 2026 plus": "CECH 2026+",
    "cech 2026+": "CECH 2026+",
    "cech 25-": "CECH 2025-",
    "cech 2025 minus": "CECH 2025-",
    "cech 2025-": "CECH 2025-",
}

check_progress_lock = threading.Lock()
check_progress = {
    "state": "idle",
    "source": "",
    "started_at": None,
    "finished_at": None,
    "total": 0,
    "pending": 0,
    "current_client": "",
    "success": 0,
    "error": 0,
    "insolvencies": 0,
    "unresolved": 0,
    "debt_relief": 0,
    "closed": 0,
    "errors": [],
    "new_documents": [],
}


def client_label(client: Client | None) -> str:
    if client is None:
        return "Neznámý klient"
    return f"{client.last_name} {client.first_name}".strip() or f"Klient #{client.id}"


def reset_check_progress(source: str) -> None:
    with check_progress_lock:
        check_progress.update(
            {
                "state": "queued",
                "source": source,
                "started_at": datetime.utcnow(),
                "finished_at": None,
                "total": 0,
                "pending": 0,
                "current_client": "",
                "success": 0,
                "error": 0,
                "insolvencies": 0,
                "unresolved": 0,
                "debt_relief": 0,
                "closed": 0,
                "errors": [],
                "new_documents": [],
            }
        )


def normalized_status_text(value: str | None) -> str:
    text = unicodedata.normalize("NFKD", str(value or "").casefold())
    return "".join(char for char in text if not unicodedata.combining(char))


def case_is_debt_relief(case: InsolvencyCase) -> bool:
    status = normalized_status_text(case.state)
    return "oddlu" in status or "splne" in status


def case_is_unresolved(case: InsolvencyCase) -> bool:
    status = normalized_status_text(case.state)
    if case_is_debt_relief(case):
        return False
    if "odskrtnuta" in status or "od krtnuta" in status:
        return False
    if "splne" in status or "zrus" in status:
        return False
    if not status or "neuveden" in status:
        return True
    return True


def case_is_closed(case: InsolvencyCase) -> bool:
    status = normalized_status_text(case.state)
    return "odskrtnuta" in status or "od krtnuta" in status or "zrus" in status


def update_check_progress(event: str, payload: dict) -> None:
    with check_progress_lock:
        if event == "started":
            total = int(payload.get("total") or 0)
            check_progress.update(
                {
                    "state": "running",
                    "total": total,
                    "pending": total,
                    "current_client": "",
                    "success": 0,
                    "error": 0,
                    "insolvencies": 0,
                    "unresolved": 0,
                    "debt_relief": 0,
                    "closed": 0,
                    "errors": [],
                    "new_documents": [],
                }
            )
            return

        if event == "client_started":
            check_progress["state"] = "running"
            check_progress["current_client"] = client_label(payload.get("client"))
            check_progress["pending"] = max(0, check_progress["pending"] - 1)
            return

        if event == "client_success":
            client = payload.get("client")
            check_progress["success"] += 1
            check_progress["current_client"] = ""
            if client is not None:
                cases = list(getattr(client, "cases", []) or [])
                check_progress["insolvencies"] += len(cases)
                check_progress["unresolved"] += sum(1 for case in cases if case_is_unresolved(case))
                check_progress["debt_relief"] += sum(1 for case in cases if case_is_debt_relief(case))
                check_progress["closed"] += sum(1 for case in cases if case_is_closed(case))
                new_document_count = int(payload.get("new_document_count") or 0)
                if (
                    not new_document_count
                    and str(check_progress.get("source") or "").startswith("Import")
                ):
                    new_document_count = int(payload.get("document_count") or 0)
                if new_document_count:
                    check_progress["new_documents"].append(
                        {
                            "client": client_label(client),
                            "client_id": getattr(client, "id", None),
                            "count": new_document_count,
                            "titles": list(payload.get("new_document_titles") or []),
                        }
                    )
            return

        if event == "client_error":
            client_name = client_label(payload.get("client"))
            error = str(payload.get("error") or "Kontrola selhala")
            check_progress["error"] += 1
            if check_progress["current_client"] != client_name:
                check_progress["pending"] = max(0, check_progress["pending"] - 1)
            check_progress["current_client"] = ""
            check_progress["errors"].append({"client": client_name, "error": error})
            return

        if event == "finished":
            check_progress["state"] = "finished"
            check_progress["finished_at"] = datetime.utcnow()
            check_progress["current_client"] = ""


def get_check_progress() -> dict:
    with check_progress_lock:
        return deepcopy(check_progress)


def dismiss_check_progress() -> None:
    with check_progress_lock:
        check_progress["state"] = "idle"
        check_progress["current_client"] = ""


def run_tracked_check(source: str, client_ids: list[int] | None = None) -> None:
    reset_check_progress(source)
    check_all_clients(progress_callback=update_check_progress, client_ids=client_ids)


def parse_claim_deadline(value: str | None):
    if not value:
        return None
    match = re.search(r"(\d{1,2})\.(\d{1,2})\.(\d{4})", str(value))
    if not match:
        return None
    day, month, year = (int(part) for part in match.groups())
    try:
        return date(year, month, day)
    except ValueError:
        return None


def add_months(value: date, months: int) -> date:
    month_index = value.month - 1 + months
    year = value.year + month_index // 12
    month = month_index % 12 + 1
    day = min(value.day, calendar.monthrange(year, month)[1])
    return date(year, month, day)


def case_started_date(case: InsolvencyCase) -> date | None:
    if case.proceeding_started_at:
        return case.proceeding_started_at.date()
    if case.started_at:
        return case.started_at
    first_filing = first_filing_date(case)
    if first_filing:
        return first_filing.date() if isinstance(first_filing, datetime) else first_filing
    return None


def effective_claim_deadline(case: InsolvencyCase | None) -> date | None:
    if case is None:
        return None
    explicit_deadline = parse_claim_deadline(case.claims_deadline)
    if explicit_deadline:
        return explicit_deadline
    started = case_started_date(case)
    if not started:
        return None
    return add_months(started, 2)


def effective_claim_deadline_label(case: InsolvencyCase | None) -> str:
    deadline = effective_claim_deadline(case)
    if not deadline:
        return "-"
    suffix = "" if case and parse_claim_deadline(case.claims_deadline) else " (vypočteno ze zahájení)"
    return f"{deadline.strftime('%d.%m.%Y')}{suffix}"


def case_study_can_be_auto_created(case: InsolvencyCase) -> bool:
    text = (case.ai_case_study or "").strip()
    if not text:
        return True
    if text.startswith("Kazuistiku se nepodařilo vytvořit"):
        if not case.ai_case_study_at:
            return True
        return datetime.utcnow() - case.ai_case_study_at > timedelta(hours=6)
    if text.startswith("AI kazuistika se připravuje") and case.ai_case_study_at:
        return datetime.utcnow() - case.ai_case_study_at > timedelta(hours=1)
    deadline = effective_claim_deadline(case)
    if deadline and datetime.now().date() > deadline and case.ai_case_study_at:
        return case.ai_case_study_at.date() <= deadline
    return False


def should_auto_create_case_study(case: InsolvencyCase, today: date | None = None) -> bool:
    today = today or datetime.now().date()
    deadline = effective_claim_deadline(case)
    if deadline and today > deadline:
        return True
    return case_is_closed(case)


def client_is_in_claim_deadline(client: Client, today: date | None = None) -> bool:
    case = primary_case(client)
    if case is None:
        return False
    deadline = effective_claim_deadline(case)
    if deadline is None:
        return False
    return (today or datetime.now().date()) <= deadline


def claim_collection_is_running(case: InsolvencyCase | None, today: date | None = None) -> bool:
    deadline = effective_claim_deadline(case)
    if deadline is None:
        return False
    return (today or datetime.now().date()) <= deadline


def claim_collection_warning() -> str:
    return "! Sběr přihlášek stále probíhá, počet přihlášek i výše pohledávek jsou průběžné údaje a nemusí být konečné."


def run_next_automatic_case_study() -> int:
    session = SessionLocal()
    case_id = None
    try:
        running_case = next(
            (
                case
                for case in session.query(InsolvencyCase).all()
                if (case.ai_case_study or "").strip().startswith("AI kazuistika se připravuje")
                and case.ai_case_study_at
                and datetime.utcnow() - case.ai_case_study_at <= timedelta(hours=1)
            ),
            None,
        )
        if running_case is not None:
            return 0

        cases = sorted(
            session.query(InsolvencyCase).all(),
            key=lambda case: (
                0 if case_is_closed(case) else 1,
                effective_claim_deadline(case) or date.max,
                case.id or 0,
            ),
        )
        for case in cases:
            if case_study_can_be_auto_created(case) and should_auto_create_case_study(case):
                case.ai_case_study_at = datetime.utcnow()
                case.ai_case_study = "AI kazuistika se připravuje na pozadí automaticky podle stavu řízení."
                case_id = case.id
                break
        session.commit()
    except Exception:
        session.rollback()
        raise
    finally:
        session.close()

    if case_id is None:
        return 0

    analyze_case_study_job(case_id)
    return 1


def ensure_automatic_case_study_job() -> None:
    session = SessionLocal()
    try:
        should_run = any(
            case_study_can_be_auto_created(case) and should_auto_create_case_study(case)
            for case in session.query(InsolvencyCase).all()
        )
    finally:
        session.close()

    if should_run:
        threading.Thread(target=run_next_automatic_case_study_locked, daemon=True).start()


def automatic_case_study_status() -> dict:
    session = SessionLocal()
    try:
        items = []
        running_seen = False
        for client in session.query(Client).order_by(Client.last_name, Client.first_name).all():
            for case in client.cases:
                if not should_auto_create_case_study(case):
                    continue
                text = (case.ai_case_study or "").strip()
                item = {
                    "client": client_label(client),
                    "case_id": case.id,
                    "state": case.state or "",
                    "deadline": case.claims_deadline or "",
                    "updated_at": case.ai_case_study_at,
                    "text": text,
                }
                if text.startswith("Kazuistiku se nepodařilo vytvořit"):
                    item["status"] = "error"
                    item["message"] = text
                elif text.startswith("AI kazuistika se připravuje"):
                    age = datetime.utcnow() - case.ai_case_study_at if case.ai_case_study_at else timedelta()
                    if case.ai_case_study_at and age > timedelta(hours=1):
                        item["status"] = "error"
                        item["message"] = "Kazuistika je zaseknutá déle než 1 hodinu, automatika ji zkusí znovu."
                    elif running_seen:
                        item["status"] = "pending"
                        item["message"] = "Čeká, až doběhne předchozí kazuistika."
                    else:
                        item["status"] = "running"
                        item["message"] = text
                        running_seen = True
                elif text:
                    item["status"] = "done"
                    item["message"] = "Hotovo"
                else:
                    item["status"] = "pending"
                    item["message"] = "Čeká na zpracování"
                items.append(item)

        return {
            "total": len(items),
            "done": sum(1 for item in items if item["status"] == "done"),
            "running": sum(1 for item in items if item["status"] == "running"),
            "pending": sum(1 for item in items if item["status"] == "pending"),
            "error": sum(1 for item in items if item["status"] == "error"),
            "errors": [item for item in items if item["status"] == "error"],
            "running_items": [item for item in items if item["status"] == "running"],
        }
    finally:
        session.close()


auto_case_study_lock = threading.Lock()


def run_next_automatic_case_study_locked() -> None:
    if not auto_case_study_lock.acquire(blocking=False):
        return
    try:
        run_next_automatic_case_study()
    finally:
        auto_case_study_lock.release()


def automatic_case_study_worker() -> None:
    time_to_wait = 20
    while True:
        threading.Event().wait(time_to_wait)
        time_to_wait = 10 * 60
        run_next_automatic_case_study_locked()


threading.Thread(target=automatic_case_study_worker, daemon=True).start()


def primary_case(client: Client):
    return client.cases[0] if client.cases else None


def clean_status(case) -> str:
    if case is None or not case.state:
        return "Bez řízení"

    status = str(case.state).strip()
    if not status or status.lower() == "stav neuveden":
        return "Bez řízení"
    return status


def status_class(case) -> str:
    status = clean_status(case).lower()
    normalized_status = unicodedata.normalize("NFKD", status)
    normalized_status = "".join(char for char in normalized_status if not unicodedata.combining(char))
    if "odskrtnuta" in normalized_status or "od krtnuta" in normalized_status:
        return "status-black"
    if "splne" in normalized_status or "oddlu" in normalized_status:
        return "status-ok"
    if "zrus" in normalized_status or "problem" in normalized_status or "selhal" in normalized_status:
        return "status-danger"
    if status == "bez řízení":
        return "status-muted"
    return "status-warning"


def event_line(case) -> str:
    if case is None:
        return "-"

    event_at = case.last_event_at
    event_description = case.last_event_description
    if not event_description:
        documents = [
            document
            for document in case.documents
            if document.event_at and document.title
        ]
        if documents:
            latest_document = max(
                documents,
                key=lambda document: (document.event_at, document.id or 0),
            )
            event_at = latest_document.event_at
            event_description = latest_document.title

    if not event_description:
        return "-"
    if event_at:
        return f"{event_at.strftime('%d.%m.%Y')} – {event_description}"
    return str(event_description)


def first_filing_date(case):
    if case is None:
        return None
    documents = [document for document in case.documents if document.event_at]
    if documents:
        first_document = min(documents, key=lambda document: document.event_at)
        return first_document.event_at
    return case.proceeding_started_at


def short_ai_summary(text: str | None) -> str:
    if not text:
        return ""
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    return " ".join(sentences[:5])


def clean_ai_value(value) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    if normalized_status_text(text) in {"-", "neuvedeno", "nelze overit", "nezjisteno"}:
        return ""
    return text


def ai_verified_data_proposal(case: InsolvencyCase | None) -> dict:
    if case is None or not case.ai_raw_result:
        return {"items": [], "has_applicable": False}
    if case.ai_category != "AI kontrola údajů z PDF":
        return {"items": [], "has_applicable": False}

    try:
        payload = json.loads(case.ai_raw_result)
    except (TypeError, ValueError):
        return {"items": [], "has_applicable": False}

    items = []
    deadline = payload.get("claims_deadline") if isinstance(payload.get("claims_deadline"), dict) else {}
    amount = payload.get("claims_total_amount") if isinstance(payload.get("claims_total_amount"), dict) else {}

    proposed_deadline = clean_ai_value(deadline.get("pdf_value"))
    if proposed_deadline:
        items.append(
            {
                "key": "claims_deadline",
                "label": "Lhůta přihlášek",
                "current": case.claims_deadline or "-",
                "proposed": proposed_deadline,
                "status": deadline.get("status") or "-",
                "confidence": deadline.get("confidence") or payload.get("confidence") or "-",
                "source": deadline.get("source") or "-",
                "will_apply": proposed_deadline != (case.claims_deadline or "").strip(),
            }
        )

    proposed_amount = clean_ai_value(amount.get("pdf_value"))
    if proposed_amount:
        items.append(
            {
                "key": "claims_total_amount",
                "label": "Výše přihlášených pohledávek",
                "current": case.claims_total_amount or "-",
                "proposed": proposed_amount,
                "status": amount.get("status") or "-",
                "confidence": payload.get("confidence") or "-",
                "source": amount.get("source") or "-",
                "will_apply": proposed_amount != (case.claims_total_amount or "").strip(),
            }
        )

    proposed_count = clean_ai_value(amount.get("claims_count"))
    if proposed_count:
        current_count = str(case.claims_count) if case.claims_count is not None else ""
        items.append(
            {
                "key": "claims_count",
                "label": "Počet přihlášek",
                "current": current_count or "-",
                "proposed": proposed_count,
                "status": amount.get("status") or "-",
                "confidence": payload.get("confidence") or "-",
                "source": amount.get("source") or "-",
                "will_apply": proposed_count != current_count,
            }
        )

    return {
        "items": items,
        "has_applicable": any(item["will_apply"] for item in items),
    }


def apply_ai_verified_data(case: InsolvencyCase) -> list[str]:
    proposal = ai_verified_data_proposal(case)
    changed = []
    values = {item["key"]: item["proposed"] for item in proposal["items"] if item["will_apply"]}

    if values.get("claims_deadline"):
        case.claims_deadline = values["claims_deadline"]
        changed.append("lhůta přihlášek")
    if values.get("claims_total_amount"):
        case.claims_total_amount = values["claims_total_amount"]
        changed.append("výše přihlášených pohledávek")
    if values.get("claims_count"):
        parsed_count = parse_optional_int(values["claims_count"])
        if parsed_count is not None:
            case.claims_count = parsed_count
            changed.append("počet přihlášek")

    return changed


def document_groups(case):
    if case is None:
        return []

    groups = {}
    for document in case.documents:
        if not document.local_path or not Path(document.local_path).exists():
            continue
        key = (
            document.event_at.strftime("%Y-%m-%d %H:%M") if document.event_at else "",
            document.title or "",
        )
        if key not in groups:
            groups[key] = {
                "id": document.id,
                "event_at": document.event_at,
                "title": document.title,
                "documents": [],
            }
        groups[key]["documents"].append(document)
        groups[key]["id"] = min(groups[key]["id"], document.id)

    return sorted(
        groups.values(),
        key=lambda group: group["event_at"] or datetime.min,
        reverse=True,
    )


def sorted_documents(case):
    if case is None:
        return []
    return sorted(
        [
            document
            for document in case.documents
            if document.local_path and Path(document.local_path).exists()
        ],
        key=lambda document: (document.event_at or datetime.min, document.id or 0),
        reverse=True,
    )


def sort_clients(clients: list[Client], sort_key: str, direction: str) -> list[Client]:
    reverse = direction == "desc"

    def key(client: Client):
        case = primary_case(client)
        if sort_key == "birth_date":
            return client.birth_date or date.min
        if sort_key == "status":
            return normalize_name(clean_status(case))
        if sort_key == "case":
            return normalize_name(case.spisova_znacka if case else "")
        if sort_key == "first_filing":
            return first_filing_date(case) or datetime.min
        return normalize_name(f"{client.last_name} {client.first_name}")

    return sorted(clients, key=key, reverse=reverse)


def status_options_for_clients(clients: list[Client]) -> list[str]:
    return sorted({clean_status(primary_case(client)) for client in clients}, key=normalize_name)


def encode_status_cookie(statuses: list[str]) -> str:
    return "||".join(quote(status, safe="") for status in statuses)


def decode_status_cookie(value: str | None) -> list[str]:
    if not value:
        return []
    return [unquote(status) for status in value.split("||") if status]


def valid_project_filters(values: list[str]) -> list[str]:
    allowed = set(PROJECT_FILTER_OPTIONS)
    result = []
    for value in values:
        normalized = normalize_project_value(value)
        if normalized in allowed and normalized not in result:
            result.append(normalized)
    return result


def normalize_project_value(value) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    normalized = normalize_name(raw)
    for option in PROJECT_FILTER_OPTIONS:
        if normalize_name(option) == normalized:
            return option
    aliased = PROJECT_IMPORT_ALIASES.get(normalized)
    if aliased:
        return aliased
    return raw


def backup_database_to_zip(zip_file: zipfile.ZipFile) -> None:
    if not DATABASE_PATH.exists():
        return

    temp_path = ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".db") as temp_file:
            temp_path = temp_file.name

        source = sqlite3.connect(str(DATABASE_PATH))
        target = sqlite3.connect(temp_path)
        try:
            source.backup(target)
        finally:
            target.close()
            source.close()

        zip_file.write(temp_path, "data/app.db")
    finally:
        if temp_path and Path(temp_path).exists():
            os.unlink(temp_path)


def add_documents_to_zip(zip_file: zipfile.ZipFile) -> None:
    documents_dir = DOCUMENTS_DIR
    if not documents_dir.exists():
        return

    for path in documents_dir.rglob("*"):
        if path.is_file():
            zip_file.write(path, Path("downloaded_documents") / path.relative_to(documents_dir))


def create_data_backup(prefix: str = "isir-data") -> Path:
    exports_dir = Path("exports")
    exports_dir.mkdir(parents=True, exist_ok=True)
    archive_path = exports_dir / f"{prefix}-{datetime.now().strftime('%Y%m%d-%H%M%S')}.zip"
    with zipfile.ZipFile(archive_path, "w", zipfile.ZIP_STORED) as zip_file:
        backup_database_to_zip(zip_file)
        add_documents_to_zip(zip_file)
        zip_file.writestr(
            "README.txt",
            "Archiv obsahuje databazi data/app.db a stazene PDF dokumenty.\n"
            "Nastaveni AI klice se z bezpecnostnich duvodu neprenasi.\n",
        )
    return archive_path


def _restore_member_relative_path(member_name: str) -> Path | None:
    normalized = member_name.replace("\\", "/").lstrip("/")
    if normalized.endswith("/"):
        return None
    marker = "/downloaded_documents/"
    if normalized.startswith("downloaded_documents/"):
        return Path(normalized).relative_to("downloaded_documents")
    if marker in normalized:
        return Path(normalized.split(marker, 1)[1])
    return None


def _safe_restore_path(base_dir: Path, relative_path: Path) -> Path:
    target = (base_dir / relative_path).resolve()
    base = base_dir.resolve()
    if base != target and base not in target.parents:
        raise ValueError("Neplatna cesta v archivu.")
    return target


def relink_document_paths() -> None:
    session = SessionLocal()
    try:
        for document in session.query(InsolvencyDocument).all():
            if not document.local_path:
                continue
            normalized = str(document.local_path).replace("\\", "/")
            marker = "/downloaded_documents/"
            if marker in normalized:
                relative = Path(normalized.split(marker, 1)[1])
                document.local_path = str(DOCUMENTS_DIR / relative)
        session.commit()
    except Exception:
        session.rollback()
        raise
    finally:
        session.close()


def restore_data_from_zip(uploaded_file) -> None:
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        archive_path = temp_path / "restore.zip"
        uploaded_file.save(archive_path)

        with zipfile.ZipFile(archive_path) as zip_file:
            names = zip_file.namelist()
            database_member = next((name for name in names if name.replace("\\", "/") == "data/app.db"), None)
            if database_member is None:
                raise ValueError("Archiv neobsahuje databazi data/app.db.")

            restored_db = temp_path / "app.db"
            with zip_file.open(database_member) as source, restored_db.open("wb") as target:
                shutil.copyfileobj(source, target)

            restored_documents = temp_path / "downloaded_documents"
            restored_documents.mkdir(parents=True, exist_ok=True)
            for member in names:
                relative_path = _restore_member_relative_path(member)
                if relative_path is None:
                    continue
                target_path = _safe_restore_path(restored_documents, relative_path)
                target_path.parent.mkdir(parents=True, exist_ok=True)
                with zip_file.open(member) as source, target_path.open("wb") as target:
                    shutil.copyfileobj(source, target)

        create_data_backup("pred-obnovou")
        engine.dispose()
        DATABASE_PATH.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(restored_db, DATABASE_PATH)

        if DOCUMENTS_DIR.exists():
            shutil.rmtree(DOCUMENTS_DIR)
        if restored_documents.exists():
            shutil.copytree(restored_documents, DOCUMENTS_DIR)
        else:
            DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)

    init_db()
    relink_document_paths()


def clear_all_client_data() -> None:
    session = SessionLocal()
    try:
        for client in session.query(Client).all():
            session.delete(client)
        session.commit()
    except Exception:
        session.rollback()
        raise
    finally:
        session.close()

    documents_dir = DOCUMENTS_DIR
    if documents_dir.exists():
        shutil.rmtree(documents_dir)
    documents_dir.mkdir(parents=True, exist_ok=True)


def recently_changed(client: Client) -> bool:
    if not client.changes:
        return False
    latest_change = client.changes[0].created_at
    if latest_change is None:
        return False
    if client.change_seen_at and client.change_seen_at >= latest_change:
        return False
    return latest_change >= datetime.utcnow() - timedelta(hours=24)


def normalize_name(value: str) -> str:
    return " ".join(value.strip().casefold().split())


def normalize_column_name(value) -> str:
    text = unicodedata.normalize("NFKD", str(value or ""))
    text = "".join(char for char in text if not unicodedata.combining(char))
    return re.sub(r"[^a-z0-9]+", " ", text.casefold()).strip()


def parse_birth_date(value):
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    text = str(value or "").strip()
    if not text:
        return None
    for date_format in ("%Y-%m-%d", "%d.%m.%Y", "%d.%m.%y", "%d/%m/%Y", "%d/%m/%y"):
        try:
            return datetime.strptime(text, date_format).date()
        except ValueError:
            pass
    return None


def parse_optional_datetime(value: str):
    text = str(value or "").strip()
    if not text:
        return None
    for date_format in ("%Y-%m-%dT%H:%M", "%Y-%m-%d", "%d.%m.%Y %H:%M", "%d.%m.%Y"):
        try:
            return datetime.strptime(text, date_format)
        except ValueError:
            pass
    return None


def parse_optional_int(value: str):
    text = str(value or "").strip()
    if not text:
        return None
    try:
        return int(text)
    except ValueError:
        return None


def first_existing_column(columns: dict[str, int], names: tuple[str, ...]) -> int | None:
    for name in names:
        if name in columns:
            return columns[name]
    return None


def find_import_column(columns: dict[str, int], kind: str) -> int | None:
    for name, index in columns.items():
        compact = name.replace(" ", "")
        if kind == "first_name" and (compact in {"jmeno", "krestnijmeno"} or compact.startswith("jm")):
            return index
        if kind == "last_name" and (compact == "prijmeni" or compact.startswith("prijmen")):
            return index
        if kind == "birth_date" and (
            compact in {"datumnarozeni", "narozeni", "datum"}
            or compact.startswith("datumnarozen")
            or compact.startswith("narozen")
        ):
            return index
    return None


def safe_folder_name(value: str) -> str:
    return re.sub(r"[^\w.-]+", "_", value, flags=re.UNICODE).strip("_")[:90] or "klient"


def safe_download_name(value: str) -> str:
    return re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", value).strip(" ._")[:120] or "soubor"


def build_case_study_docx(case: InsolvencyCase) -> BytesIO:
    client = case.client
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Heading 1"].font.name = "Arial"

    document.add_heading("Kazuistika AI", level=0)

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    metadata = [
        ("Klient", f"{client.last_name} {client.first_name}" if client else "-"),
        ("Datum narození", format_date(client.birth_date) if client else "-"),
        ("Spisová značka", case.spisova_znacka or "-"),
        ("Stav řízení", clean_status(case)),
        ("Vytvořeno", format_datetime(case.ai_case_study_at)),
    ]
    if effective_claim_deadline(case):
        metadata.append(("Lhůta přihlášek", effective_claim_deadline_label(case)))
    collection_running = claim_collection_is_running(case)
    if case.claims_count:
        value = str(case.claims_count)
        if collection_running:
            value = f"{value} ! sběr stále probíhá"
        metadata.append(("Počet přihlášek", value))
    if case.claims_total_amount:
        value = case.claims_total_amount
        if collection_running:
            value = f"{value}\n{claim_collection_warning()}"
        metadata.append(("Výše pohledávek", value))

    for label, value in metadata:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True

    document.add_paragraph()
    for block in (case.ai_case_study or "").split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.endswith(":") and "\n" not in block:
            document.add_heading(block.rstrip(":"), level=1)
            continue
        lines = block.splitlines()
        if len(lines) > 1 and lines[0].endswith(":"):
            document.add_heading(lines[0].rstrip(":"), level=1)
            for line in lines[1:]:
                line = line.strip()
                if line.startswith("- "):
                    document.add_paragraph(line[2:], style="List Bullet")
                elif line:
                    document.add_paragraph(line)
        else:
            document.add_paragraph(block)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def delete_client_files(client: Client) -> None:
    folder_name = safe_folder_name(
        f"{client.last_name}_{client.first_name}_{client.birth_date.isoformat() if client.birth_date else 'bez_data'}"
    )
    client_path = DOCUMENTS_DIR / folder_name
    if client_path.exists():
        shutil.rmtree(client_path)

    case_ids = [case.id for case in client.cases if case.id is not None]
    for case_id in case_ids:
        path = DOCUMENTS_DIR / str(case_id)
        if path.exists():
            shutil.rmtree(path)


app.jinja_env.globals.update(
    primary_case=primary_case,
    clean_status=clean_status,
    status_class=status_class,
    event_line=event_line,
    first_filing_date=first_filing_date,
    short_ai_summary=short_ai_summary,
    document_groups=document_groups,
    sorted_documents=sorted_documents,
    ai_verified_data_proposal=ai_verified_data_proposal,
    effective_claim_deadline_label=effective_claim_deadline_label,
    claim_collection_is_running=claim_collection_is_running,
    claim_collection_warning=claim_collection_warning,
    has_gemini_api_key=has_gemini_api_key,
    recently_changed=recently_changed,
)


@app.template_filter("format_datetime")
def format_datetime(value: datetime | None) -> str:
    if value is None:
        return "-"
    return value.strftime("%d.%m.%Y %H:%M")


@app.template_filter("datetime_local")
def datetime_local(value: datetime | None) -> str:
    if value is None:
        return ""
    return value.strftime("%Y-%m-%dT%H:%M")


@app.template_filter("format_date")
def format_date(value) -> str:
    if value is None:
        return "-"
    return value.strftime("%d.%m.%Y")


@app.get("/")
def index():
    ensure_automatic_case_study_job()
    error = request.args.get("error", "")
    message = request.args.get("message", "")
    sort_key = request.args.get("sort") or request.cookies.get("index_sort", "client")
    direction = request.args.get("direction") or request.cookies.get("index_direction", "asc")
    if request.args.get("clear_status") == "1":
        selected_statuses = []
    elif "status" in request.args:
        selected_statuses = [value for value in request.args.getlist("status") if value]
    else:
        selected_statuses = decode_status_cookie(request.cookies.get("index_statuses"))
    if request.args.get("clear_project") == "1":
        selected_projects = []
    elif "project_filter" in request.args or "project" in request.args:
        selected_projects = valid_project_filters([value for value in request.args.getlist("project") if value])
    else:
        selected_projects = valid_project_filters(decode_status_cookie(request.cookies.get("index_projects")))
    if request.args.get("clear_status") == "1":
        selected_in_deadline = False
    elif "deadline_filter" in request.args:
        selected_in_deadline = request.args.get("in_deadline") == "1"
    else:
        selected_in_deadline = request.cookies.get("index_in_deadline") == "1"
    allowed_sorts = {"client", "birth_date", "status", "case", "first_filing"}
    if sort_key not in allowed_sorts:
        sort_key = "client"
    if direction not in {"asc", "desc"}:
        direction = "asc"

    def next_direction(column: str) -> str:
        if sort_key == column and direction == "asc":
            return "desc"
        return "asc"

    def sort_url(column: str) -> str:
        return url_for(
            "index",
            sort=column,
            direction=next_direction(column),
            status=selected_statuses,
            project=selected_projects,
            in_deadline="1" if selected_in_deadline else None,
            deadline_filter="1" if selected_in_deadline else None,
        )

    session = SessionLocal()
    try:
        all_clients = session.query(Client).order_by(Client.last_name, Client.first_name).all()
        status_options = status_options_for_clients(all_clients)
        important_statuses = [status for status in status_options if status != "Bez řízení"]
        clients = all_clients
        if selected_projects:
            clients = [
                client
                for client in clients
                if normalize_project_value(client.project) in selected_projects
            ]
        if selected_statuses:
            filtered_clients = [
                client
                for client in clients
                if clean_status(primary_case(client)) in selected_statuses
            ]
            if filtered_clients:
                clients = filtered_clients
            else:
                selected_statuses = []
        if selected_in_deadline:
            clients = [client for client in clients if client_is_in_claim_deadline(client)]
        clients = sort_clients(clients, sort_key, direction)
        response = make_response(
            render_template(
                "index.html",
                clients=clients,
                error=error,
                message=message,
                sort_key=sort_key,
                direction=direction,
                selected_statuses=selected_statuses,
                selected_projects=selected_projects,
                selected_in_deadline=selected_in_deadline,
                status_options=status_options,
                important_statuses=important_statuses,
                project_filter_options=PROJECT_FILTER_OPTIONS,
                automatic_case_study_status=automatic_case_study_status(),
                next_direction=next_direction,
                sort_url=sort_url,
                check_progress=get_check_progress(),
            )
        )
        response.set_cookie("index_sort", sort_key, max_age=60 * 60 * 24 * 365)
        response.set_cookie("index_direction", direction, max_age=60 * 60 * 24 * 365)
        response.set_cookie(
            "index_statuses",
            encode_status_cookie(selected_statuses),
            max_age=60 * 60 * 24 * 365,
        )
        response.set_cookie(
            "index_projects",
            encode_status_cookie(selected_projects),
            max_age=60 * 60 * 24 * 365,
        )
        response.set_cookie(
            "index_in_deadline",
            "1" if selected_in_deadline else "0",
            max_age=60 * 60 * 24 * 365,
        )
        return response
    finally:
        session.close()


@app.get("/data/export")
def export_data():
    archive_path = create_data_backup()

    return send_file(
        archive_path,
        mimetype="application/zip",
        as_attachment=True,
        download_name=archive_path.name,
    )


@app.post("/data/import")
def import_data():
    uploaded_file = request.files.get("backup_file")
    if uploaded_file is None or not uploaded_file.filename.lower().endswith(".zip"):
        return redirect(url_for("index", error="Nahrajte ZIP zálohu ze Stažení dat."))

    try:
        restore_data_from_zip(uploaded_file)
    except Exception as exc:
        return redirect(url_for("index", error=f"Obnova dat se nepodařila: {exc}"))

    return redirect(url_for("index", message="Data byla obnovena ze zálohy."))


@app.post("/check-progress/dismiss")
def dismiss_check_progress_route():
    dismiss_check_progress()
    return redirect(url_for("index"))


@app.post("/data/clear")
def clear_data():
    clear_all_client_data()
    return redirect(url_for("index", message="Všechna klientská data byla vymazána."))


@app.get("/settings")
def settings():
    saved = request.args.get("saved") == "1"
    return render_template("settings.html", saved=saved, has_key=has_gemini_api_key())


@app.post("/settings")
def save_settings_route():
    api_key = request.form.get("gemini_api_key", "").strip()
    if api_key:
        set_gemini_api_key(api_key)
    return redirect(url_for("settings", saved="1"))


@app.post("/clients")
def add_client():
    first_name = request.form.get("first_name", "").strip()
    last_name = request.form.get("last_name", "").strip()
    birth_date_raw = request.form.get("birth_date", "").strip()

    if not first_name or not last_name or not birth_date_raw:
        return redirect(url_for("index"))

    session = SessionLocal()
    try:
        birth_date = datetime.strptime(birth_date_raw, "%Y-%m-%d").date()
        existing_clients = session.query(Client).all()
        duplicate = next(
            (
                client
                for client in existing_clients
                if normalize_name(client.first_name) == normalize_name(first_name)
                and normalize_name(client.last_name) == normalize_name(last_name)
                and client.birth_date == birth_date
            ),
            None,
        )
        if duplicate is not None:
            return redirect(url_for("index", error="Klient už je v seznamu."))

        client = Client(
            first_name=first_name,
            last_name=last_name,
            birth_date=birth_date,
        )
        session.add(client)
        try:
            session.commit()
        except IntegrityError:
            session.rollback()
            return redirect(url_for("index", error="Klient už je v seznamu."))
    finally:
        session.close()

    return redirect(url_for("index"))


@app.post("/clients/import")
def import_clients():
    uploaded_file = request.files.get("clients_file")
    if uploaded_file is None or not uploaded_file.filename.lower().endswith(".xlsx"):
        return redirect(url_for("index", error="Nahrajte soubor XLSX."))

    try:
        workbook = load_workbook(uploaded_file, read_only=True, data_only=True)
        sheet = workbook.active
        rows = sheet.iter_rows(values_only=True)
        header = next(rows, None)
    except Exception:
        return redirect(url_for("index", error="Soubor XLSX se nepodařilo načíst."))

    if not header:
        return redirect(url_for("index", error="Soubor XLSX je prázdný."))

    columns = {normalize_column_name(name): index for index, name in enumerate(header)}
    first_name_index = first_existing_column(columns, ("jmeno", "krestni jmeno"))
    last_name_index = first_existing_column(columns, ("prijmeni",))
    birth_date_index = first_existing_column(columns, ("datum narozeni", "narozeni", "datum"))
    project_index = first_existing_column(columns, ("projekt", "project"))
    if first_name_index is None:
        first_name_index = find_import_column(columns, "first_name")
    if last_name_index is None:
        last_name_index = find_import_column(columns, "last_name")
    if birth_date_index is None:
        birth_date_index = find_import_column(columns, "birth_date")

    if first_name_index is None or last_name_index is None or birth_date_index is None:
        return redirect(url_for("index", error="XLSX musí obsahovat sloupce: jméno, příjmení, datum narození. Sloupec projekt je volitelný."))

    imported = 0
    skipped = 0
    invalid = 0
    imported_client_ids = []
    session = SessionLocal()
    try:
        existing_clients = session.query(Client).all()
        existing_by_key = {
            (normalize_name(client.first_name), normalize_name(client.last_name), client.birth_date): client
            for client in existing_clients
        }

        for row in rows:
            first_name = str(row[first_name_index] or "").strip() if first_name_index < len(row) else ""
            last_name = str(row[last_name_index] or "").strip() if last_name_index < len(row) else ""
            birth_date_value = row[birth_date_index] if birth_date_index < len(row) else None
            birth_date = parse_birth_date(birth_date_value)
            project = normalize_project_value(row[project_index]) if project_index is not None and project_index < len(row) else ""

            if not first_name or not last_name or birth_date is None:
                invalid += 1
                continue

            key = (normalize_name(first_name), normalize_name(last_name), birth_date)
            existing_client = existing_by_key.get(key)
            if existing_client is not None:
                if project:
                    existing_client.project = project
                skipped += 1
                continue

            client = Client(first_name=first_name, last_name=last_name, birth_date=birth_date, project=project or None)
            session.add(client)
            session.flush()
            existing_by_key[key] = client
            imported_client_ids.append(client.id)
            imported += 1

        session.commit()
    except Exception:
        session.rollback()
        return redirect(url_for("index", error="Import se nepodařilo uložit."))
    finally:
        session.close()

    if imported:
        scheduler.add_job(
            run_tracked_check,
            args=["Import klientů", imported_client_ids],
            id=f"import_isir_check_{int(datetime.utcnow().timestamp())}",
            replace_existing=False,
        )

    message = f"Importováno {imported} klientů. Duplicit přeskočeno {skipped}. Neplatných řádků {invalid}."
    if imported:
        message += " Kontrola ISIR běží na pozadí."
    return redirect(url_for("index", message=message))


@app.post("/clients/<int:client_id>/delete")
def delete_client(client_id: int):
    session = SessionLocal()
    try:
        client = session.get(Client, client_id)
        if client is not None:
            delete_client_files(client)
            session.delete(client)
            session.commit()
    finally:
        session.close()

    return redirect(url_for("index"))


@app.post("/clients/<int:client_id>/update")
def update_client(client_id: int):
    first_name = request.form.get("first_name", "").strip()
    last_name = request.form.get("last_name", "").strip()
    birth_date_raw = request.form.get("birth_date", "").strip()

    if not first_name or not last_name or not birth_date_raw:
        return redirect(url_for("client_detail", client_id=client_id))

    session = SessionLocal()
    try:
        client = session.get(Client, client_id)
        if client is None:
            return redirect(url_for("index"))

        birth_date = datetime.strptime(birth_date_raw, "%Y-%m-%d").date()
        duplicate = next(
            (
                other_client
                for other_client in session.query(Client).filter(Client.id != client_id).all()
                if normalize_name(other_client.first_name) == normalize_name(first_name)
                and normalize_name(other_client.last_name) == normalize_name(last_name)
                and other_client.birth_date == birth_date
            ),
            None,
        )
        if duplicate is not None:
            return redirect(url_for("client_detail", client_id=client_id))

        old_folder = safe_folder_name(
            f"{client.last_name}_{client.first_name}_{client.birth_date.isoformat() if client.birth_date else 'bez_data'}"
        )
        client.first_name = first_name
        client.last_name = last_name
        client.birth_date = birth_date

        case = primary_case(client)
        if case is not None:
            case.state = request.form.get("state", "").strip() or None
            debtor_name = request.form.get("debtor_name", request.form.get("court", "")).strip()
            case.debtor_name = debtor_name or None
            case.spisova_znacka = request.form.get("spisova_znacka", "").strip() or case.spisova_znacka
            case.proceeding_started_at = parse_optional_datetime(request.form.get("proceeding_started_at", ""))
            case.last_event_at = parse_optional_datetime(request.form.get("last_event_at", ""))
            case.last_event_description = request.form.get("last_event_description", "").strip() or None
            case.document_count = parse_optional_int(request.form.get("document_count", "")) or 0
            case.claims_deadline = request.form.get("claims_deadline", "").strip() or None
            case.claims_count = parse_optional_int(request.form.get("claims_count", ""))
            case.claims_total_amount = request.form.get("claims_total_amount", "").strip() or None

        client.last_checked_at = parse_optional_datetime(request.form.get("last_checked_at", ""))

        new_folder = safe_folder_name(f"{client.last_name}_{client.first_name}_{client.birth_date.isoformat()}")
        old_path = DOCUMENTS_DIR / old_folder
        new_path = DOCUMENTS_DIR / new_folder
        if old_path.exists() and old_path != new_path and not new_path.exists():
            old_path.rename(new_path)
            for case in client.cases:
                for document in case.documents:
                    if document.local_path:
                        document.local_path = str(Path(str(document.local_path).replace(str(old_path), str(new_path))))
        session.commit()
    except ValueError:
        session.rollback()
    except IntegrityError:
        session.rollback()
    finally:
        session.close()

    return redirect(url_for("client_detail", client_id=client_id))


@app.get("/clients/<int:client_id>")
def client_detail(client_id: int):
    session = SessionLocal()
    try:
        client = session.get(Client, client_id)
        if client is None:
            return redirect(url_for("index"))
        if recently_changed(client):
            client.change_seen_at = datetime.utcnow()
            session.commit()
        return render_template("client_detail.html", client=client)
    finally:
        session.close()


@app.get("/documents/<int:document_id>")
def open_document(document_id: int):
    session = SessionLocal()
    try:
        document = session.get(InsolvencyDocument, document_id)
        if document is None or not document.local_path:
            abort(404)
        path = Path(document.local_path)
        if not path.exists():
            abort(404)
        return send_file(path, mimetype="application/pdf", download_name=path.name)
    finally:
        session.close()


@app.get("/documents/<int:document_id>/download")
def download_document(document_id: int):
    session = SessionLocal()
    try:
        document = session.get(InsolvencyDocument, document_id)
        if document is None or not document.local_path:
            abort(404)
        path = Path(document.local_path)
        if not path.exists():
            abort(404)
        return send_file(path, mimetype="application/pdf", as_attachment=True, download_name=path.name)
    finally:
        session.close()


@app.post("/documents/<int:document_id>/delete")
def delete_document(document_id: int):
    next_url = request.form.get("next") or url_for("index")
    session = SessionLocal()
    try:
        document = session.get(InsolvencyDocument, document_id)
        if document is not None:
            path = Path(document.local_path) if document.local_path else None
            if path and path.exists():
                try:
                    path.unlink()
                except OSError:
                    pass
            case = document.case
            document.local_path = None
            document.file_size = None
            if case is not None and case.document_count:
                case.document_count = sum(
                    1
                    for item in case.documents
                    if item.local_path and Path(item.local_path).exists()
                )
            session.commit()
    finally:
        session.close()
    return redirect(next_url)


@app.get("/cases/<int:case_id>/case-study/download")
def download_case_study(case_id: int):
    session = SessionLocal()
    try:
        case = session.get(InsolvencyCase, case_id)
        if case is None or not case.ai_case_study:
            abort(404)

        client = case.client
        data = build_case_study_docx(case)
        filename = safe_download_name(
            f"kazuistika_{client.last_name}_{client.first_name}_{case.spisova_znacka}.docx"
            if client
            else f"kazuistika_{case.id}.docx"
        )
        return send_file(
            data,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename,
        )
    finally:
        session.close()


@app.route("/check-now", methods=["GET", "POST"])
def check_now():
    if request.method == "GET":
        return redirect(url_for("index"))

    scheduler.add_job(
        run_tracked_check,
        args=["Ruční kontrola"],
        id=f"manual_isir_check_{int(datetime.utcnow().timestamp())}",
        replace_existing=False,
    )
    return redirect(url_for("index", message="Kontrola ISIR běží na pozadí."))


@app.post("/check-projects")
def check_selected_projects():
    selected_projects = valid_project_filters([value for value in request.form.getlist("project") if value])
    if not selected_projects:
        return redirect(url_for("index", error="Vyberte alespoň jeden projekt pro kontrolu ISIR."))

    session = SessionLocal()
    try:
        clients = session.query(Client).all()
        client_ids = [
            client.id
            for client in clients
            if normalize_project_value(client.project) in selected_projects
        ]
    finally:
        session.close()

    if not client_ids:
        return redirect(url_for("index", project_filter=1, project=selected_projects, error="Ve vybraných projektech není žádný klient."))

    project_label = ", ".join(selected_projects)
    scheduler.add_job(
        run_tracked_check,
        args=[f"Kontrola projektů: {project_label}", client_ids],
        id=f"project_isir_check_{int(datetime.utcnow().timestamp())}",
        replace_existing=False,
    )
    return redirect(
        url_for(
            "index",
            project_filter=1,
            project=selected_projects,
            message=f"Kontrola ISIR pro vybrané projekty běží na pozadí ({len(client_ids)} klientů).",
        )
    )


@app.post("/cases/<int:case_id>/analyze")
def analyze_case(case_id: int):
    next_url = request.form.get("next") or url_for("index")
    session = SessionLocal()
    try:
        case = session.get(InsolvencyCase, case_id)
        if case is not None:
            case.ai_checked_at = datetime.utcnow()
            case.ai_category = "Analýza běží"
            case.ai_summary = "AI shrnutí se připravuje na pozadí. Stránka se obnoví sama, případně můžete obnovit ručně (F5)."
            session.commit()
            scheduler.add_job(
                analyze_case_latest_document_job,
                args=[case_id],
                id=f"ai_analysis_{case_id}_{int(datetime.utcnow().timestamp())}",
                replace_existing=False,
            )
    finally:
        session.close()

    return redirect(next_url)


@app.post("/cases/<int:case_id>/case-study")
def create_case_study(case_id: int):
    next_url = request.form.get("next") or url_for("index")
    session = SessionLocal()
    try:
        case = session.get(InsolvencyCase, case_id)
        if case is not None:
            case.ai_case_study_at = datetime.utcnow()
            case.ai_case_study = "AI kazuistika se připravuje na pozadí. Stránka se obnoví sama, případně můžete obnovit ručně (F5)."
            session.commit()
            scheduler.add_job(
                analyze_case_study_job,
                args=[case_id],
                id=f"ai_case_study_{case_id}_{int(datetime.utcnow().timestamp())}",
                replace_existing=False,
            )
    finally:
        session.close()

    return redirect(next_url)


@app.post("/cases/<int:case_id>/verify-data")
def verify_case_data(case_id: int):
    next_url = request.form.get("next") or url_for("index")
    session = SessionLocal()
    try:
        case = session.get(InsolvencyCase, case_id)
        if case is not None:
            case.ai_checked_at = datetime.utcnow()
            case.ai_category = "AI kontrola údajů běží"
            case.ai_summary = "AI ověření údajů z PDF se připravuje na pozadí. Stránka se obnoví sama, případně můžete obnovit ručně (F5)."
            session.commit()
            scheduler.add_job(
                analyze_case_data_verification_job,
                args=[case_id],
                id=f"ai_data_verification_{case_id}_{int(datetime.utcnow().timestamp())}",
                replace_existing=False,
            )
    finally:
        session.close()

    return redirect(next_url)


@app.post("/cases/<int:case_id>/apply-verified-data")
def apply_verified_case_data(case_id: int):
    next_url = request.form.get("next") or url_for("index")
    session = SessionLocal()
    try:
        case = session.get(InsolvencyCase, case_id)
        if case is not None:
            changed = apply_ai_verified_data(case)
            if changed:
                case.ai_checked_at = datetime.utcnow()
                case.ai_category = "AI ověřené údaje zapsány"
                case.ai_summary = "Zapsáno po potvrzení: " + ", ".join(changed) + "."
            session.commit()
    finally:
        session.close()

    return redirect(next_url)


if __name__ == "__main__":
    from waitress import serve

    serve(app, host="127.0.0.1", port=5000, threads=8)
